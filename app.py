import streamlit as st
from audio_recorder_streamlit import audio_recorder
import whisper
import tempfile
import os
import io
from docx import Document
import google.generativeai as genai

# 1. Cấu hình trang
st.set_page_config(page_title="AI Smart Note", page_icon="🎙️", layout="wide")

# --- PHẦN CSS TÙY CHỈNH ---
st.markdown("""
    <style>
        /* Nhập Google Fonts */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600&family=Space+Grotesk:wght@700&display=swap');

        /* Áp dụng font cho toàn bộ ứng dụng */
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
        }

        /* Tùy chỉnh Tiêu đề (H1) */
        h1 {
            font-family: 'Space Grotesk', sans-serif;
            text-align: center;
            background: -webkit-linear-gradient(45deg, #FF4B4B, #FF914D);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-size: 3.5rem !important;
            padding-bottom: 20px;
        }

        /* Tùy chỉnh Tiêu đề phụ (H2, H3) */
        h2, h3 {
            font-family: 'Space Grotesk', sans-serif;
            color: #333;
        }

        /* Khu vực ghi âm (Recording Box) */
        .recording-box {
            background-color: #f8f9fa; /* Màu nền xám nhẹ */
            border: 2px dashed #d1d5db; /* Viền nét đứt */
            border-radius: 20px;
            padding: 30px;
            text-align: center;
            margin-bottom: 30px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        }
        
        /* Làm đẹp nút tải xuống (stButton) */
        .stButton > button {
            width: 100%;
            border-radius: 10px;
            font-weight: 600;
            padding: 0.5rem 1rem;
            border: none;
            transition: all 0.3s ease;
        }
        
        /* Hiệu ứng khi di chuột vào nút */
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }

        /* Căn giữa các thông báo st.info, st.success */
        .stAlert {
            border-radius: 10px;
        }
    </style>
""", unsafe_allow_html=True)
# --------------------------

# 2. Xử lý API Key từ Secrets
api_key = None
if "GOOGLE_API_KEY" in st.secrets:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
else:
    st.error("⚠️ Vui lòng cấu hình GOOGLE_API_KEY trong Secrets.")
    st.stop()

# 3. Khởi tạo Session State
if 'notes' not in st.session_state:
    st.session_state.notes = []

# 4. Tải Whisper
@st.cache_resource
def load_whisper_model():
    return whisper.load_model("base")

# Hiển thị tiêu đề
st.title("🎙️ AI Smart Note")
st.markdown("<p style='text-align: center; color: #666; margin-top: -20px; margin-bottom: 40px;'>Biến giọng nói thành văn bản và tóm tắt thông minh</p>", unsafe_allow_html=True)

# Tải model (ẩn spinner để giao diện đẹp hơn, chỉ hiện khi cần)
model = load_whisper_model()

# --- LOGIC XỬ LÝ ---
def transcribe_audio(audio_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
        temp_audio.write(audio_bytes)
        temp_filename = temp_audio.name
    try:
        result = model.transcribe(temp_filename, language="vi")
        return result["text"]
    except Exception as e:
        return f"Lỗi: {e}"
    finally:
        if os.path.exists(temp_filename): os.remove(temp_filename)

def summarize_text(text):
    try:
        model_gemini = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"Sửa lỗi chính tả, tóm tắt ý chính và liệt kê hành động từ văn bản sau: '{text}'"
        response = model_gemini.generate_content(prompt)
        return response.text
    except Exception as e:
        return str(e)

def create_docx(original, summary):
    doc = Document()
    doc.add_heading('Biên bản AI', 0)
    doc.add_heading('Tóm tắt', 1)
    doc.add_paragraph(summary)
    doc.add_heading('Chi tiết', 1)
    doc.add_paragraph(original)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- GIAO DIỆN CHÍNH ---

# Tạo bố cục 3 cột để căn giữa nút ghi âm
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    # Tạo container giả lập hiệu ứng thẻ bài (Card)
    st.markdown('<div class="recording-box">', unsafe_allow_html=True)
    st.write("👇 **Nhấn vào icon bên dưới để bắt đầu**")
    
    # Nút ghi âm
    audio_bytes = audio_recorder(
        text="",
        recording_color="#ff4b4b", # Màu đỏ khi ghi
        neutral_color="#FF914D",   # Màu cam khi chờ
        icon_name="microphone",
        icon_size="4x",            # Icon lớn hơn
    )
    st.markdown('</div>', unsafe_allow_html=True)

# Xử lý khi có âm thanh
if audio_bytes:
    # Hiển thị thanh phát lại nhỏ gọn ở giữa
    col1_a, col2_a, col3_a = st.columns([1, 2, 1])
    with col2_a:
        st.audio(audio_bytes, format="audio/wav")

    with st.status("🤖 AI đang làm việc...", expanded=True) as status:
        st.write("👂 Đang nghe và gỡ băng (Whisper)...")
        transcript = transcribe_audio(audio_bytes)
        
        st.write("🧠 Đang suy nghĩ và tóm tắt (Gemini)...")
        summary = summarize_text(transcript)
        
        status.update(label="✅ Hoàn tất!", state="complete", expanded=False)

    # Hiển thị kết quả
    if transcript and summary:
        st.divider()
        
        # Layout kết quả
        c1, c2 = st.columns(2)
        
        with c1:
            st.markdown("### 📝 Tóm tắt & Hành động")
            st.info(summary)
        
        with c2:
            st.markdown("### 📄 Văn bản gốc")
            with st.container(height=300): # Thanh cuộn cho văn bản dài
                st.write(transcript)
        
        # Nút tải về nằm giữa
        st.markdown("<br>", unsafe_allow_html=True)
        col_d1, col_d2, col_d3 = st.columns([1, 1, 1])
        with col_d2:
            docx = create_docx(transcript, summary)
            st.download_button(
                label="📥 TẢI VỀ FILE WORD",
                data=docx,
                file_name="SmartNote_AI.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True # Nút rộng full container
            )

        # Lưu lịch sử
        note_data = {"original": transcript, "summary": summary}
        if not st.session_state.notes or st.session_state.notes[-1]["original"] != transcript:
            st.session_state.notes.append(note_data)

# Lịch sử (Footer)
if st.session_state.notes:
    st.divider()
    st.caption("Lịch sử phiên làm việc gần đây:")
    for note in reversed(st.session_state.notes[-3:]): # Chỉ hiện 3 cái mới nhất
        with st.expander(f"Note: {note['original'][:50]}..."):
            st.write(note['summary'])
